#!/usr/bin/env python3
"""
Hebrew PDF OCR Interactive Editor – Tesseract / Surya (GPU) Dual Engine (Fixed)
=============================================================================

This version fixes:
- Unterminated string literal (`txt.count('\n')`) bug.
- OCR results & tables indentation (they now sit inside `if uploaded_pdf:` block).
- Removed duplicate per‑shape OCR bookkeeping lines.
- Replaced outdated Surya imports with supported `surya.models.load_predictors` API (0.14.6).
- Added robust Surya text extraction + debug option.
- Prevents duplicate shapes (hash filtering & zero-size filter).
- Cursor alignment & zoom/pan mapping stable.

Dependencies (example):
```
pip install streamlit==1.31.1 streamlit-drawable-canvas==0.9.3 pytesseract pillow opencv-python-headless pymupdf python-docx pandas
sudo apt-get install tesseract-ocr tesseract-ocr-heb
# Optional GPU OCR (Surya):
pip install surya-ocr torch --extra-index-url https://download.pytorch.org/whl/cu121
```
"""
from __future__ import annotations
import os, io, json, math, hashlib, tempfile, zipfile, time
from dataclasses import dataclass, field
from typing import List, Tuple, Dict, Any, Optional

import streamlit as st
from streamlit_drawable_canvas import st_canvas
from PIL import Image, ImageDraw
import numpy as np
import cv2
import pytesseract
import fitz  # PyMuPDF
import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

st.set_page_config(page_title="Hebrew PDF OCR", layout="wide")

# ------------------------- CONFIG ------------------------- #
DEFAULT_BASE_DPI = 220
MAX_ZOOM = 2.5
MIN_ZOOM = 0.35
ZOOM_STEP = 0.10
PAN_STEP_RATIO = 0.06
LINE_MERGE_TOL = 6
MIN_CELL_SIZE = 8
HEB_RANGE = (0x0590, 0x05FF)
RTL_RATIO = 0.35
TARGET_VIEWPORT_WIDTH = 1100
MAX_VIEWPORT_HEIGHT = 1500
HOUGH_MAX_LINES = 220
SURYA_LANG = 'heb'
DEBUG_SURYA = False  # set True to dump raw surya structures once

# ------------------------- STATE INIT ------------------------- #
STATE_DEFAULTS: Dict[str, Any] = {
    'pdf_pages': None,
    'pdf_name': None,
    'BASE_DPI': DEFAULT_BASE_DPI,
    'zoom': {},
    'pan': {},
    'shapes': {},
    'ocr': {},
    'shape_counter': 0,
    'last_ocr_hashes': {},
    'messages': [],
    'selected_shape': {},
    'surya_ctx': None,
    'surya_error': None,
    'stats': {'regions_ocr_time': []},
    'search_term': '',
    'pending_rerun': False,
    'pp_mode': 'adaptive',
    'surya_debug_dumped': False,
}
for k,v in STATE_DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ------------------------- UTILITIES ------------------------- #

def st_safe_rerun():
    if not st.session_state.pending_rerun:
        st.session_state.pending_rerun = True
        st.rerun()

def sanitize_text(t: str) -> str:
    if not t: return ''
    return t.replace('‏','').replace('‎','').replace('```','‵‵‵').strip()

def hebrew_ratio(text: str) -> float:
    if not text: return 0.0
    letters = ''.join(ch for ch in text if ch.isalpha())
    if not letters: return 0.0
    heb = sum(1 for ch in letters if HEB_RANGE[0] <= ord(ch) <= HEB_RANGE[1])
    return heb / len(letters)

def looks_hebrew(text: str) -> bool:
    return hebrew_ratio(text) >= RTL_RATIO

# ------------------------- DATA MODEL ------------------------- #
@dataclass
class Shape:
    id: str
    type: str                # rect | polygon | line
    bbox: Optional[Tuple[int,int,int,int]] = None
    points: Optional[List[Tuple[float,float]]] = None
    text: Optional[str] = None
    label: str = ""
    engine: str = ""         # 'tesseract' | 'surya'
    hash: str = field(default="")
    last_psm: Optional[int] = None

    def geometry_hash(self) -> str:
        if self.type == 'rect' and self.bbox:
            raw = f"rect:{self.bbox}"
        elif self.points:
            raw = f"{self.type}:{[(round(x,1),round(y,1)) for x,y in self.points]}"
        else:
            raw = self.id
        self.hash = hashlib.md5(raw.encode()).hexdigest()
        return self.hash

# ------------------------- PDF & IMAGE UTILS ------------------------- #

def load_pdf_pages(pdf_bytes: bytes, dpi: int) -> List[Image.Image]:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(pdf_bytes); path = tmp.name
    doc = fitz.open(path)
    pages = []
    for pg in doc:
        pix = pg.get_pixmap(dpi=dpi)
        pages.append(Image.open(io.BytesIO(pix.tobytes("png"))))
    doc.close(); os.unlink(path)
    return pages

def pil_to_cv(im: Image.Image) -> np.ndarray:
    return cv2.cvtColor(np.array(im), cv2.COLOR_RGB2BGR)

def cv_to_pil(arr: np.ndarray) -> Image.Image:
    return Image.fromarray(cv2.cvtColor(arr, cv2.COLOR_BGR2RGB))

# ------------------------- PREPROCESS ------------------------- #

def preprocess_for_ocr(pil_img: Image.Image, enhance_numbers: bool = False, mode: str = "adaptive") -> Image.Image:
    if mode == 'none':
        return pil_img
    cv_img = pil_to_cv(pil_img)
    if mode == 'mild':
        gray = cv2.cvtColor(cv_img, cv2.COLOR_BGR2GRAY)
        gray = cv2.equalizeHist(gray)
        if enhance_numbers:
            kernel = cv2.getStructuringElement(cv2.MORPH_RECT,(2,2))
            gray = cv2.dilate(gray,kernel,1)
        return cv_to_pil(cv2.cvtColor(gray, cv2.COLOR_GRAY2BGR))
    # adaptive
    gray = cv2.cvtColor(cv_img, cv2.COLOR_BGR2GRAY)
    gray = cv2.medianBlur(gray,3)
    th = cv2.adaptiveThreshold(gray,255,cv2.ADAPTIVE_THRESH_GAUSSIAN_C,cv2.THRESH_BINARY,31,5)
    if enhance_numbers:
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT,(2,2))
        th = cv2.dilate(th,kernel,1)
    return cv_to_pil(cv2.cvtColor(th, cv2.COLOR_GRAY2BGR))

# ------------------------- CROPPING ------------------------- #

def polygon_mask_crop(img: Image.Image, pts: List[Tuple[float,float]]) -> Image.Image:
    cv_img = pil_to_cv(img)
    mask = np.zeros(cv_img.shape[:2], dtype=np.uint8)
    ipts = np.array([[int(x),int(y)] for x,y in pts], dtype=np.int32)
    cv2.fillPoly(mask,[ipts],255)
    fg = cv2.bitwise_and(cv_img,cv_img,mask=mask)
    white = np.full_like(cv_img,255)
    inv = cv2.bitwise_not(mask)
    merged = cv2.add(fg, cv2.bitwise_and(white,white,mask=inv))
    ys,xs = np.where(mask>0)
    if len(xs)==0: return cv_to_pil(merged)
    x0,x1=xs.min(),xs.max(); y0,y1=ys.min(),ys.max()
    return cv_to_pil(merged[y0:y1+1, x0:x1+1])

def crop_shape(shape: Shape, base_img: Image.Image) -> Optional[Image.Image]:
    if shape.type=='rect' and shape.bbox:
        return base_img.crop(shape.bbox)
    if shape.type=='polygon' and shape.points:
        return polygon_mask_crop(base_img, shape.points)
    return None

# ------------------------- TESSERACT SUPPORT ------------------------- #
@st.cache_data(show_spinner=False)
def detect_supported_oems(langs: str) -> List[int]:
    im = Image.new("RGB",(60,30),"white"); d=ImageDraw.Draw(im); d.text((2,2),"אב12", fill=(0,0,0))
    out=[]
    for oem in (3,1,2,0):
        try:
            pytesseract.image_to_string(im, lang=langs, config=f"--psm 7 --oem {oem}")
            out.append(oem)
        except pytesseract.TesseractError:
            continue
    return out or [3]

if 'supported_oems' not in st.session_state:
    st.session_state.supported_oems = detect_supported_oems('heb+eng')
    if st.session_state.supported_oems == [3]:
        st.session_state.messages.append("Only Tesseract OEM 3 detected (no legacy components).")

# ------------------------- SURYA (Modern API) ------------------------- #
@st.cache_resource(show_spinner=True)
def init_surya():
    try:
        import torch
        from surya.models import load_predictors
        device = 'cuda' if torch.cuda.is_available() else 'cpu'
        preds = load_predictors(device=device)
        det = preds.get('detection'); rec = preds.get('recognition')
        if not det or not rec:
            return None, "Missing detection or recognition predictor."
        return {'device':device,'det':det,'rec':rec}, None
    except Exception as e:
        return None, str(e)

def ensure_surya():
    if st.session_state.surya_ctx or st.session_state.surya_error:
        return
    ctx, err = init_surya()
    if err:
        st.session_state.surya_error = err
        st.session_state.messages.append(f"Surya init failed: {err}")
    else:
        st.session_state.surya_ctx = ctx

def run_surya_ocr(img: Image.Image) -> str:
    ensure_surya()
    ctx = st.session_state.surya_ctx
    if not ctx:
        return "[Surya unavailable]"
    det = ctx['det']; rec = ctx['rec']
    try:
        arr = np.array(img.convert('RGB'))
        det_out = det([arr])
        boxes=[]
        entry = det_out[0] if isinstance(det_out,(list,tuple)) else det_out
        if isinstance(entry, dict):
            if entry.get('polygons'):
                for poly in entry['polygons']:
                    xs=[p[0] for p in poly]; ys=[p[1] for p in poly]
                    boxes.append([min(xs),min(ys),max(xs),max(ys)])
            if not boxes and entry.get('boxes'):
                for b in entry['boxes']:
                    if hasattr(b,'tolist'): b=b.tolist()
                    if len(b)>=4: boxes.append(b[:4])
        if not boxes:
            boxes=[[0,0,arr.shape[1],arr.shape[0]]]
        lines=[]
        debug_once = DEBUG_SURYA and not st.session_state.surya_debug_dumped
        for (x0,y0,x1,y1) in boxes:
            x0=int(max(0,x0)); y0=int(max(0,y0))
            x1=int(min(arr.shape[1],x1)); y1=int(min(arr.shape[0],y1))
            if x1-x0<4 or y1-y0<4: continue
            crop = arr[y0:y1, x0:x1]
            rec_out = rec([crop])
            if debug_once:
                st.write("Surya det sample:", det_out)
                st.write("Surya rec sample:", rec_out)
                st.session_state.surya_debug_dumped = True
            val=None
            if isinstance(rec_out,(list,tuple)) and rec_out:
                cand=rec_out[0]
                if isinstance(cand,dict):
                    for k in ('text','value','prediction'):
                        if k in cand: val=cand[k]; break
                elif isinstance(cand,str):
                    val=cand
            elif isinstance(rec_out,dict):
                val = rec_out.get('text') or rec_out.get('value')
            if isinstance(val,(list,tuple)):
                val=''.join(map(str,val))
            if val and val.strip():
                lines.append(val.strip())
        return sanitize_text('\n'.join(lines)) or "[Surya empty]"
    except Exception as e:
        return f"[Surya OCR error] {e}".strip()

# ------------------------- OCR (Tesseract) ------------------------- #

def resolve_psm(psm_choice: Any, shape: Shape) -> int:
    if psm_choice == 'Auto':
        if shape.type=='rect' and shape.bbox:
            x0,y0,x1,y1 = shape.bbox
            w=x1-x0; h=y1-y0
            if w*h < 2500 or w/(h+1e-3) > 6:
                return 13
            return 6
        return 6
    return int(psm_choice)

def run_ocr_tesseract(img: Image.Image, allow_eng: bool, numeric_bias: bool, psm: int, oem_eff: int, pp_mode: str, fallback: bool=True) -> str:
    langs = 'heb' + ('+eng' if allow_eng else '')
    cfg = f"--psm {psm} --oem {oem_eff}"
    if numeric_bias:
        cfg += ' -c tessedit_char_whitelist=0123456789₪.,:%/-+()'
    pre = preprocess_for_ocr(img, enhance_numbers=numeric_bias, mode=pp_mode)
    try:
        txt = pytesseract.image_to_string(pre, lang=langs, config=cfg).strip()
    except pytesseract.TesseractError as e:
        return f"[Tesseract error] {e}".strip()
    if fallback and (len(txt.strip())==0 or txt.count('\n') < 1):
        for alt in ("mild","none"):
            pre2 = preprocess_for_ocr(img, enhance_numbers=numeric_bias, mode=alt)
            try:
                alt_txt = pytesseract.image_to_string(pre2, lang=langs, config=cfg).strip()
                if len(alt_txt) > len(txt):
                    txt = alt_txt
                    break
            except pytesseract.TesseractError:
                continue
    return txt

def run_ocr_on_shape(shape: Shape, base_img: Image.Image, engine: str, allow_eng: bool, numeric_bias: bool, psm_choice: Any, oem: int):
    if shape.type not in ('rect','polygon'): return
    crop = crop_shape(shape, base_img)
    if crop is None: return
    start = time.time()
    if engine=='surya':
        txt = run_surya_ocr(crop)
    else:
        psm_val = resolve_psm(psm_choice, shape)
        oem_eff = oem if oem in st.session_state.supported_oems else 3
        pp_mode = st.session_state.get('pp_mode','adaptive')
        txt = run_ocr_tesseract(crop, allow_eng, numeric_bias, psm_val, oem_eff, pp_mode)
        shape.last_psm = psm_val
    shape.text = sanitize_text(txt)
    shape.engine = engine
    shape.geometry_hash()
    st.session_state.stats['regions_ocr_time'].append(time.time()-start)

# ------------------------- TABLE HELPERS ------------------------- #

def group_lines(line_shapes: List[Shape], tol: int = LINE_MERGE_TOL) -> Tuple[List[int], List[int]]:
    h,v=[],[]
    for sh in line_shapes:
        if not sh.points or len(sh.points)<2: continue
        xs=[p[0] for p in sh.points]; ys=[p[1] for p in sh.points]
        dx=max(xs)-min(xs); dy=max(ys)-min(ys)
        if dy<=tol and dx>dy: h.append(int(round(sum(ys)/len(ys))))
        elif dx<=tol and dy>dx: v.append(int(round(sum(xs)/len(xs))))
    def merge(vals):
        vals=sorted(vals); out=[]
        for val in vals:
            if not out or abs(val-out[-1])>tol: out.append(val)
        return out
    return merge(h), merge(v)

def build_table(base_img: Image.Image, h_lines: List[int], v_lines: List[int], engine: str, allow_eng: bool, numeric_bias: bool, psm_choice: Any, oem: int) -> Optional[List[List[str]]]:
    if len(h_lines)<2 or len(v_lines)<2: return None
    w,h = base_img.size
    if 0 not in h_lines: h_lines=[0]+h_lines
    if h not in h_lines: h_lines=h_lines+[h]
    if 0 not in v_lines: v_lines=[0]+v_lines
    if w not in v_lines: v_lines=v_lines+[w]
    h_lines=sorted(set(h_lines)); v_lines=sorted(set(v_lines))
    rows=[]
    for ri in range(len(h_lines)-1):
        y0,y1=h_lines[ri],h_lines[ri+1]
        if y1-y0 < MIN_CELL_SIZE: continue
        row=[]
        for ci in range(len(v_lines)-1):
            x0,x1=v_lines[ci],v_lines[ci+1]
            if x1-x0 < MIN_CELL_SIZE: continue
            cell=base_img.crop((x0,y0,x1,y1))
            if engine=='surya':
                txt=run_surya_ocr(cell)
            else:
                dummy=Shape('tmp','rect',bbox=(x0,y0,x1,y1))
                psm_val=resolve_psm(psm_choice,dummy)
                oem_eff = oem if oem in st.session_state.supported_oems else 3
                txt=run_ocr_tesseract(cell,allow_eng,numeric_bias,psm_val,oem_eff,st.session_state.pp_mode)
            row.append(sanitize_text(txt))
        if row: rows.append(row)
    return rows or None

# ------------------------- ZOOM / PAN ------------------------- #

def get_zoom(page:int)->float: return st.session_state.zoom.get(page,1.0)

def set_zoom(page:int,val:float): st.session_state.zoom[page]=max(MIN_ZOOM,min(MAX_ZOOM,val))

def get_pan(page:int)->Tuple[int,int]: return st.session_state.pan.get(page,(0,0))

def set_pan(page:int,dx:int,dy:int): st.session_state.pan[page]=(dx,dy)

def make_zoom_canvas(page:int, base: Image.Image, zoom: float, pan: Tuple[int,int]):
    w,h = base.size
    if page not in st.session_state.zoom and w > TARGET_VIEWPORT_WIDTH:
        zoom = min(1.0, TARGET_VIEWPORT_WIDTH / w)
    if zoom * w > TARGET_VIEWPORT_WIDTH * 1.05:
        zoom = TARGET_VIEWPORT_WIDTH / w
        set_zoom(page, zoom)
    zw,zh = int(w*zoom), int(h*zoom)
    scaled = base.resize((zw,zh), Image.BICUBIC)
    canvas_w = min(TARGET_VIEWPORT_WIDTH, zw)
    canvas_h = min(max(zh, int(canvas_w*(h/w))), MAX_VIEWPORT_HEIGHT)
    cx = (canvas_w - zw)//2
    cy = (canvas_h - zh)//2
    dx,dy = pan
    if zw <= canvas_w: dx=0
    else: dx = max(canvas_w-zw, min(0, dx))
    if zh <= canvas_h: dy=0
    else: dy = max(canvas_h-zh, min(0, dy))
    set_pan(page, dx, dy)
    canvas = Image.new('RGB',(canvas_w,canvas_h),'white')
    canvas.paste(scaled,(cx+dx, cy+dy))
    return canvas, (dx,dy), (cx,cy), zoom

# ------------------------- SHAPE MGMT ------------------------- #

def next_shape_id()->str:
    st.session_state.shape_counter += 1
    return f"S{st.session_state.shape_counter}"

def add_shape(page:int, shape:Shape):
    shape.geometry_hash()
    st.session_state.shapes.setdefault(page,[]).append(shape)

def delete_shape(page:int, shape_id:str):
    st.session_state.shapes[page] = [s for s in st.session_state.shapes.get(page,[]) if s.id!=shape_id]

def move_shape(page:int, shape_id:str, direction:int):
    arr=st.session_state.shapes.get(page,[])
    for i,s in enumerate(arr):
        if s.id==shape_id:
            j=max(0,min(len(arr)-1,i+direction))
            if i!=j: arr[i],arr[j]=arr[j],arr[i]
            return

# ------------------------- EXPORT ------------------------- #

def export_docx() -> io.BytesIO:
    doc = Document()
    for pg, shapes in st.session_state.shapes.items():
        doc.add_heading(f"Page {pg+1}", level=2)
        for s in shapes:
            if s.type in ('rect','polygon') and s.text:
                para = doc.add_paragraph(s.text)
                if looks_hebrew(s.text):
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    para.paragraph_format.right_to_left = True
        for tbl in st.session_state.ocr.get(pg,{}).get('tables',[]):
            cells=tbl['cells']
            if not cells: continue
            rows=len(cells); cols=max(len(r) for r in cells)
            t = doc.add_table(rows=rows, cols=cols)
            for r_i,row in enumerate(cells):
                for c_i,cell_text in enumerate(row):
                    p=t.rows[r_i].cells[c_i].paragraphs[0]
                    p.text = cell_text
                    if looks_hebrew(cell_text):
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        p.paragraph_format.right_to_left = True
    bio=io.BytesIO(); doc.save(bio); bio.seek(0); return bio

def export_json_zip(pages: List[Image.Image]) -> io.BytesIO:
    bundle={'meta':{'pages':len(pages),'dpi':st.session_state.BASE_DPI},'pages':[]}
    for pg,shapes in st.session_state.shapes.items():
        entry={'page_index':pg,'shapes':[],'tables':st.session_state.ocr.get(pg,{}).get('tables',[])}
        for s in shapes:
            entry['shapes'].append({'id':s.id,'type':s.type,'bbox':s.bbox,'points':s.points,'text':s.text,'engine':s.engine,'label':s.label,'last_psm':s.last_psm})
        bundle['pages'].append(entry)
    bio=io.BytesIO()
    with zipfile.ZipFile(bio,'w',zipfile.ZIP_DEFLATED) as zf:
        zf.writestr('bundle.json', json.dumps(bundle, ensure_ascii=False, indent=2))
        for pg_entry in bundle['pages']:
            for tbl in pg_entry['tables']:
                rows=tbl['cells']
                if not rows: continue
                df=pd.DataFrame(rows)
                zf.writestr(f"tables/page{pg_entry['page_index']+1}_{tbl['id']}.csv", df.to_csv(index=False))
    bio.seek(0); return bio

# ------------------------- SIDEBAR ------------------------- #
st.title("📄 Hebrew PDF OCR (Tesseract / Surya)")
with st.sidebar:
    st.header("PDF & Render")
    uploaded_pdf = st.file_uploader("PDF (Hebrew / Mixed)", type=['pdf'])
    new_dpi = st.number_input("Base DPI",120,400,st.session_state.BASE_DPI,20)
    if uploaded_pdf and new_dpi != st.session_state.BASE_DPI:
        st.session_state.BASE_DPI = new_dpi
        st.session_state.pdf_pages = load_pdf_pages(uploaded_pdf.read(), dpi=new_dpi)
        st.session_state.pdf_name = uploaded_pdf.name
        for key in ('zoom','pan','shapes','ocr','last_ocr_hashes'):
            st.session_state[key].clear()

    st.header("OCR Settings")
    engine_choice_label = st.radio("Engine", ["Tesseract","Surya (GPU)"] , index=0)
    engine_choice = 'surya' if engine_choice_label.startswith('Surya') else 'tesseract'
    allow_eng = st.checkbox("Detect English", True)
    numeric_bias = st.checkbox("Numeric bias (digits/currency)", True)
    psm_choice = st.selectbox("PSM", ["Auto",6,13,11,3,4,5,12], index=0)
    pp_mode = st.selectbox("Preprocess Mode", ["adaptive","mild","none"], index=0)
    st.session_state.pp_mode = pp_mode
    oem_opts = [o for o in st.session_state.supported_oems if o in (3,1)] or [3]
    oem_choice = st.selectbox("Tesseract OEM", oem_opts, index=0)
    auto_line_detect = st.checkbox("Auto-detect lines (Hough)", False)

    st.header("Search")
    st.session_state.search_term = st.text_input("Term", st.session_state.search_term)
    show_only_matches = st.checkbox("Only matching regions", False)

    batch_all = st.button("Batch OCR (new/changed)")
    st.caption("If Surya unavailable: install surya-ocr & proper torch, then rerun.")

# ------------------------- MAIN ------------------------- #
if uploaded_pdf:
    if st.session_state.pdf_name != uploaded_pdf.name or st.session_state.pdf_pages is None:
        st.session_state.pdf_pages = load_pdf_pages(uploaded_pdf.read(), dpi=st.session_state.BASE_DPI)
        st.session_state.pdf_name = uploaded_pdf.name
        for key in ('zoom','pan','shapes','ocr','last_ocr_hashes'):
            st.session_state[key].clear()

    pages = st.session_state.pdf_pages
    total = len(pages)
    nav_col,_ = st.columns([1,5])
    with nav_col:
        page_num = st.number_input("Page",1,total,1)
        st.caption(f"{page_num}/{total}")
    page_index = page_num - 1
    base_img = pages[page_index]

    if page_index not in st.session_state.zoom:
        if base_img.width > TARGET_VIEWPORT_WIDTH:
            st.session_state.zoom[page_index] = min(1.0, TARGET_VIEWPORT_WIDTH / base_img.width)
        else:
            st.session_state.zoom[page_index] = 1.0
        st.session_state.pan[page_index] = (0,0)

    z = get_zoom(page_index); pan = get_pan(page_index)
    zoom_canvas, applied_pan, center_off, z = make_zoom_canvas(page_index, base_img, z, pan)

    st.subheader("Page Preview")
    interaction = st.radio("Mode", ["Pan/Drag","Add Rectangle","Add Polygon","Add Line","Edit"], horizontal=True)
    draw_mode = {"Pan/Drag":"transform","Add Rectangle":"rect","Add Polygon":"polygon","Add Line":"line","Edit":"transform"}[interaction]

    def serialize_shapes(shapes: List[Shape]) -> List[Dict[str,Any]]:
        objs=[]; dx,dy = applied_pan; cx,cy = center_off
        for s in shapes:
            if s.type=='rect' and s.bbox:
                x0,y0,x1,y1 = s.bbox
                objs.append({"type":"rect","left":x0*z+dx+cx,"top":y0*z+dy+cy,
                             "width":(x1-x0)*z,"height":(y1-y0)*z,
                             "stroke":"#0044FF","fill":"rgba(0,150,255,0.25)","shape_id":s.id})
            elif s.points:
                path=[]
                for i,(x,y) in enumerate(s.points):
                    path.append(["M" if i==0 else "L", x*z+dx+cx, y*z+dy+cy])
                if s.type=='polygon': path.append(["Z"])
                objs.append({"type":"path","path":path,
                             "stroke":"#FF0055" if s.type=='line' else "#0044FF",
                             "fill":"transparent" if s.type=='line' else "rgba(0,150,255,0.25)",
                             "shape_id":s.id})
        return objs

    existing_shapes = st.session_state.shapes.get(page_index,[])
    initial_objs = serialize_shapes(existing_shapes)

    try:
        canvas = st_canvas(
            fill_color="rgba(0,150,255,0.25)",
            stroke_width=2,
            stroke_color="#FF0055" if draw_mode=='line' else "#0044FF",
            background_image=zoom_canvas,
            update_streamlit=True,
            height=zoom_canvas.height,
            width=zoom_canvas.width,
            drawing_mode=draw_mode,
            initial_drawing={"version":1,"objects":initial_objs},
            key=f"canvas_{page_index}"
        )
    except Exception as e:
        st.error(f"Canvas error: {e}")
        canvas=None

    # View Controls
    vc = st.columns(8)
    if vc[0].button("Zoom -"): set_zoom(page_index,z-ZOOM_STEP); st_safe_rerun()
    if vc[1].button("Zoom +"): set_zoom(page_index,z+ZOOM_STEP); st_safe_rerun()
    if vc[2].button("Pan Left"): dx,dy=get_pan(page_index); set_pan(page_index, dx-int(base_img.width*PAN_STEP_RATIO), dy); st_safe_rerun()
    if vc[3].button("Pan Right"): dx,dy=get_pan(page_index); set_pan(page_index, dx+int(base_img.width*PAN_STEP_RATIO), dy); st_safe_rerun()
    if vc[4].button("Pan Up"): dx,dy=get_pan(page_index); set_pan(page_index, dx, dy-int(base_img.height*PAN_STEP_RATIO)); st_safe_rerun()
    if vc[5].button("Pan Down"): dx,dy=get_pan(page_index); set_pan(page_index, dx, dy+int(base_img.height*PAN_STEP_RATIO)); st_safe_rerun()
    if vc[6].button("Fit Width"): set_zoom(page_index, min(MAX_ZOOM, max(MIN_ZOOM, TARGET_VIEWPORT_WIDTH / base_img.width))); st_safe_rerun()
    if vc[7].button("Reset View"): set_zoom(page_index,1.0); set_pan(page_index,0,0); st_safe_rerun()

    # Ingest new shapes
    if canvas and canvas.json_data:
        objs = canvas.json_data.get('objects', [])
        new_objs = [o for o in objs if not o.get('shape_id')]
        if draw_mode in ('rect','polygon','line') and new_objs:
            dx,dy = applied_pan; cx,cy = center_off
            existing_hashes = {s.geometry_hash() for s in existing_shapes}
            for obj in new_objs:
                t = obj.get('type')
                if t=='rect':
                    left = int((obj['left'] - dx - cx)/z)
                    top  = int((obj['top']  - dy - cy)/z)
                    w_box = int(obj['width']/z); h_box = int(obj['height']/z)
                    if w_box>4 and h_box>4:
                        bbox=(left,top,left+w_box,top+h_box)
                        probe=Shape('tmp','rect',bbox=bbox)
                        gh=probe.geometry_hash()
                        if gh not in existing_hashes:
                            add_shape(page_index, Shape(id=next_shape_id(), type='rect', bbox=bbox, label='Rect'))
                            existing_hashes.add(gh)
                elif t=='path' and obj.get('path'):
                    pts=[]
                    for seg in obj['path']:
                        if isinstance(seg,list) and len(seg)>=3:
                            pts.append(((seg[1]-dx-cx)/z, (seg[2]-dy-cy)/z))
                    if len(pts)>=2:
                        xs=[p[0] for p in pts]; ys=[p[1] for p in pts]
                        dx_span=max(xs)-min(xs); dy_span=max(ys)-min(ys)
                        shape_type='line'
                        if len(pts)>=3 and (dx_span*dy_span)>50:
                            shape_type='polygon'
                        probe=Shape('tmp',shape_type,points=pts)
                        gh=probe.geometry_hash()
                        if gh not in existing_hashes:
                            add_shape(page_index, Shape(id=next_shape_id(), type=shape_type, points=pts, label='Poly' if shape_type=='polygon' else 'Line'))
                            existing_hashes.add(gh)
            st_safe_rerun()
        # Edit updates
        if draw_mode=='transform' and not interaction.startswith('Add'):
            with_ids=[o for o in objs if o.get('shape_id')]
            if len(with_ids)==len(existing_shapes):
                dx,dy=applied_pan; cx,cy=center_off
                for obj in with_ids:
                    sid=obj.get('shape_id'); shape = next((s for s in existing_shapes if s.id==sid),None)
                    if not shape: continue
                    if obj['type']=='rect' and shape.type=='rect':
                        left=int((obj['left']-dx-cx)/z); top=int((obj['top']-dy-cy)/z)
                        w_box=int(obj['width']/z); h_box=int(obj['height']/z)
                        if w_box>4 and h_box>4:
                            shape.bbox=(left,top,left+w_box,top+h_box); shape.geometry_hash()
                    elif obj['type']=='path' and shape.points:
                        pts=[]
                        for seg in obj.get('path',[]):
                            if isinstance(seg,list) and len(seg)>=3:
                                pts.append(((seg[1]-dx-cx)/z,(seg[2]-dy-cy)/z))
                        if len(pts)>=2:
                            shape.points=pts; shape.geometry_hash()

    # Shape Manager
    st.markdown("### ✏️ Shapes")
    shapes = st.session_state.shapes.get(page_index,[])
    if shapes:
        for s in shapes:
            cols = st.columns([2,1,1,1,1,1])
            sel = cols[0].checkbox(s.id, value=(st.session_state.selected_shape.get(page_index)==s.id), key=f"sel_{s.id}")
            if sel:
                st.session_state.selected_shape[page_index]=s.id
            elif st.session_state.selected_shape.get(page_index)==s.id:
                st.session_state.selected_shape[page_index]=None
            new_label = cols[0].text_input("Label", value=s.label or s.id, key=f"lbl_{s.id}")
            if new_label!=s.label: s.label=new_label
            if cols[1].button("▲", key=f"up_{s.id}"): move_shape(page_index,s.id,-1); st_safe_rerun()
            if cols[2].button("▼", key=f"down_{s.id}"): move_shape(page_index,s.id,1); st_safe_rerun()
            if cols[3].button("OCR", key=f"ocr_{s.id}"):
                run_ocr_on_shape(s, base_img, engine_choice, allow_eng, numeric_bias, psm_choice, oem_choice); st_safe_rerun()
            if cols[4].button("Alt", key=f"alt_{s.id}"):
                alt = 'tesseract' if s.engine=='surya' else 'surya'
                run_ocr_on_shape(s, base_img, alt, allow_eng, numeric_bias, psm_choice, oem_choice); st_safe_rerun()
            if cols[5].button("Del", key=f"del_{s.id}"):
                delete_shape(page_index,s.id); st_safe_rerun()
        ca1,ca2,ca3 = st.columns(3)
        if ca1.button("OCR New/Changed"):
            hashes = st.session_state.last_ocr_hashes.setdefault(page_index,set())
            for s in shapes:
                if s.type in ('rect','polygon'):
                    gh=s.geometry_hash()
                    if gh not in hashes or not s.text:
                        run_ocr_on_shape(s, base_img, engine_choice, allow_eng, numeric_bias, psm_choice, oem_choice)
                        hashes.add(gh)
            st.success("OCR complete.")
        if ca2.button("Build Table From Lines"):
            line_shapes=[sh for sh in shapes if sh.type=='line']
            if auto_line_detect:
                cv_img = pil_to_cv(base_img)
                edges = cv2.Canny(cv_img,50,150,apertureSize=3)
                hough = cv2.HoughLinesP(edges,1,np.pi/180,threshold=120,minLineLength=60,maxLineGap=5)
                if hough is not None:
                    for i,(x1,y1,x2,y2) in enumerate(hough[:,0][:HOUGH_MAX_LINES]):
                        line_shapes.append(Shape(id=f"H{i}_{x1}_{y1}", type='line', points=[(x1,y1),(x2,y2)], label='H'))
            h_lines,v_lines = group_lines(line_shapes)
            tbl_cells = build_table(base_img,h_lines,v_lines,engine_choice,allow_eng,numeric_bias,psm_choice,oem_choice)
            if tbl_cells:
                st.session_state.ocr.setdefault(page_index,{}).setdefault('tables',[]).append({
                    'id':f"T{len(st.session_state.ocr.get(page_index,{}).get('tables',[]))+1}",
                    'cells':tbl_cells,'h':h_lines,'v':v_lines,'engine':engine_choice
                })
                st.success("Table added.")
            else:
                st.warning("Need ≥2 horizontal & vertical lines.")
        if ca3.button("Clear Shapes"):
            st.session_state.shapes[page_index]=[]; st.success("Shapes cleared."); st_safe_rerun()
    else:
        st.caption("No shapes yet – draw some.")

    # Batch OCR across pages
    if batch_all:
        all_shapes=[(pg,s) for pg,sl in st.session_state.shapes.items() for s in sl if s.type in ('rect','polygon')]
        total_shapes=len(all_shapes) or 1
        progress = st.progress(0.0)
        for i,(pg,s) in enumerate(all_shapes,1):
            img=st.session_state.pdf_pages[pg]
            gh=s.geometry_hash()
            seen=st.session_state.last_ocr_hashes.setdefault(pg,set())
            if gh not in seen or not s.text:
                run_ocr_on_shape(s,img,engine_choice,allow_eng,numeric_bias,psm_choice,oem_choice)
                seen.add(gh)
            progress.progress(i/total_shapes)
        st.success("Batch OCR complete.")

    # OCR RESULTS
    st.markdown("### 🔎 OCR Results")
    term = (st.session_state.search_term or '').lower()
    for s in shapes:
        if s.type in ('rect','polygon') and s.text is not None:
            if term and term not in s.text.lower() and show_only_matches:
                continue
            rtl = looks_hebrew(s.text)
            display = sanitize_text(s.text)
            engine_tag = s.engine or '—'
            if rtl:
                st.markdown(
                    "<div style='border:1px solid #ddd;padding:6px;border-radius:4px;background:#fcfcfc;margin-bottom:4px'>"
                    f"<b>{s.label or s.id}</b> ({engine_tag}) • RTL<br><div dir='rtl' style='white-space:pre-wrap;font-family:monospace;'>{display}</div></div>",
                    unsafe_allow_html=True)
            else:
                st.markdown(f"**{s.label or s.id}** ({engine_tag})")
                st.code(display, language='text')

    page_tables = st.session_state.ocr.get(page_index,{}).get('tables',[])
    for tbl in page_tables:
        st.markdown(f"**Table {tbl['id']}** ({tbl['engine']})")
        max_cols = max(len(r) for r in tbl['cells']) if tbl['cells'] else 0
        normalized = [row + [""]*(max_cols-len(row)) for row in tbl['cells']]
        df = pd.DataFrame(normalized)
        st.dataframe(df, use_container_width=True)
        st.download_button(f"CSV {tbl['id']}", data=df.to_csv(index=False).encode('utf-8'), file_name=f"{tbl['id']}.csv")

    # EXPORTS
    st.markdown("### 💾 Export")
    e1,e2 = st.columns(2)
    if e1.button("Export DOCX"):
        bio = export_docx()
        st.download_button("Download DOCX", data=bio.getvalue(), file_name="hebrew_ocr_output.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    if e2.button("Export JSON/ZIP"):
        bio2 = export_json_zip(pages)
        st.download_button("Download ZIP", data=bio2.getvalue(), file_name="hebrew_ocr_bundle.zip", mime="application/zip")
else:
    st.info("Upload a PDF to begin.")

# ------------------------- DIAGNOSTICS ------------------------- #
with st.expander("🔧 Diagnostics"):
    st.write("Supported Tesseract OEMs:", st.session_state.get('supported_oems'))
    st.write("Surya available:", st.session_state.surya_ctx is not None)
    if st.session_state.surya_error:
        st.write("Surya error:", st.session_state.surya_error)
    times = st.session_state.stats['regions_ocr_time']
    if times:
        st.write(f"OCR regions: {len(times)} | Avg time: {sum(times)/len(times):.2f}s")

for msg in st.session_state.messages:
    st.warning(msg)
st.session_state.messages.clear()
st.session_state.pending_rerun = False
